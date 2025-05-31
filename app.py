
# NRTNMED: Nursing Document Extractor & Converter (with Templates and Multi-platform Support)

import requests
from bs4 import BeautifulSoup
import re
from fpdf import FPDF
from docx import Document
import streamlit as st
import os

class NRTNMEDConverter:
    def __init__(self, url, course, institution, objectives, references):
        self.url = url
        self.title = "Untitled"
        self.slides = []
        self.course = course
        self.institution = institution
        self.objectives = objectives
        self.references = references

    def extract_slideshare_text(self):
        try:
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(self.url, headers=headers)
            soup = BeautifulSoup(response.content, 'html.parser')
            self.title = soup.title.string if soup.title else "SlideShare Document"
            for script in soup.find_all("script"):
                if 'slide_text' in script.text:
                    matches = re.findall(r'"text":"(.*?)"', script.text)
                    self.slides.extend([re.sub(r'\\n', '\n', m) for m in matches])
            return True
        except Exception as e:
            st.error(f"SlideShare extraction failed: {e}")
            return False

    def extract_scribd_text(self):
        try:
            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(self.url, headers=headers)
            soup = BeautifulSoup(response.content, 'html.parser')
            self.title = soup.title.string if soup.title else "Scribd Document"
            paragraphs = soup.find_all('p')
            for p in paragraphs:
                text = p.get_text(strip=True)
                if text:
                    self.slides.append(text)
            return True
        except Exception as e:
            st.error(f"Scribd extraction failed: {e}")
            return False

    def extract_slideplayer_text(self):
        try:
            response = requests.get(self.url)
            soup = BeautifulSoup(response.text, 'html.parser')
            self.title = soup.title.string if soup.title else "SlidePlayer Document"
            slides = soup.select(".slide-content")
            for s in slides:
                self.slides.append(s.get_text(strip=True))
            return len(self.slides) > 0
        except Exception as e:
            st.error(f"SlidePlayer extraction failed: {e}")
            return False

    def extract_slidesgo_text(self):
        try:
            response = requests.get(self.url)
            soup = BeautifulSoup(response.text, 'html.parser')
            self.title = soup.title.string if soup.title else "SlidesGo Document"
            slides = soup.find_all("section")
            for s in slides:
                text = s.get_text(strip=True)
                if len(text) > 20:
                    self.slides.append(text)
            return len(self.slides) > 0
        except Exception as e:
            st.error(f"SlidesGo extraction failed: {e}")
            return False

    def add_nursing_template_slides(self):
        cover = f"{self.title}\nCourse: {self.course}\nInstitution: {self.institution}"
        self.slides.insert(0, cover)

        if self.objectives:
            self.slides.insert(1, "Objectives:\n" + self.objectives)
        if self.references:
            self.slides.append("References:\n" + self.references)

    def save_as_pdf(self, filename="NRTNMED_Output.pdf"):
        self.add_nursing_template_slides()
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt=self.title, ln=True, align='C')
        pdf.ln(10)
        for i, slide in enumerate(self.slides):
            pdf.multi_cell(0, 10, f"Slide {i + 1}:
{slide}
")
        pdf.output(filename)
        return filename

    def save_as_docx(self, filename="NRTNMED_Output.docx"):
        self.add_nursing_template_slides()
        doc = Document()
        doc.add_heading(self.title, level=1)
        for i, slide in enumerate(self.slides):
            doc.add_heading(f"Slide {i + 1}", level=2)
            doc.add_paragraph(slide)
        doc.save(filename)
        return filename

# === STREAMLIT INTERFACE ===
st.title("📘 NRTNMED: Nursing Document Converter")
st.markdown("Supports SlideShare, Scribd, SlidePlayer, SlidesGo. Adds nursing templates: cover, objectives, references.")

url = st.text_input("Paste your document link (SlideShare, Scribd, SlidePlayer, SlidesGo):")
course = st.text_input("Course Name (e.g., M.Sc. Nursing, Post Basic B.Sc.)")
institution = st.text_input("Institution Name")
objectives = st.text_area("Enter Objectives (optional)")
references = st.text_area("Enter References (optional)")

export_pdf = st.checkbox("Export as PDF", value=True)
export_docx = st.checkbox("Export as DOCX", value=True)

if st.button("Convert Document") and url:
    converter = NRTNMEDConverter(url, course, institution, objectives, references)
    success = False

    if "slideshare.net" in url:
        success = converter.extract_slideshare_text()
    elif "scribd.com" in url:
        success = converter.extract_scribd_text()
    elif "slideplayer.com" in url:
        success = converter.extract_slideplayer_text()
    elif "slidesgo.com" in url:
        success = converter.extract_slidesgo_text()
    else:
        st.warning("Only SlideShare, Scribd, SlidePlayer, and SlidesGo links are supported.")

    if success:
        st.success("✅ Content extracted successfully!")
        if export_pdf:
            pdf_file = converter.save_as_pdf()
            with open(pdf_file, "rb") as f:
                st.download_button("📥 Download PDF", f, file_name=pdf_file)
        if export_docx:
            docx_file = converter.save_as_docx()
            with open(docx_file, "rb") as f:
                st.download_button("📥 Download DOCX", f, file_name=docx_file)
    else:
        st.error("❌ Failed to extract content. Ensure the document is public and supported.")
